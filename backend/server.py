from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer
from docx.enum.text import WD_BREAK
import requests
import io
import os

app = Flask(__name__)
CORS(app)

@app.route('/ping', methods=['GET', 'HEAD'])
def ping():
    return jsonify({"status": "ok"}), 200

@app.route('/merge-first-syllabus', methods=['POST'])
def merge_first_syllabus():
    data = request.json
    title = data['title']
    syllabus_url = data['syllabusUrl']
    
    # Fetch the syllabus file from the provided URL
    response = requests.get(syllabus_url)
    if not response.ok:
        return {"error": f"Failed to fetch syllabus: {response.status_code}"}, 500
    
    syllabus_bytes = response.content
    
    # Load the syllabus document
    doc = Document(io.BytesIO(syllabus_bytes))
    
    # Insert a new paragraph at the very beginning
    first_paragraph = doc.paragraphs[0]
    # Create a new paragraph before the first one
    new_paragraph = first_paragraph.insert_paragraph_before(title)
    new_paragraph.runs[0].font.name = 'Cambria'
    new_paragraph.runs[0].font.size = Pt(11)
    new_paragraph.runs[0].font.bold=True  # Slightly larger to look like a title
    
    # Save the modified document to a buffer
    merged_buffer = io.BytesIO()
    doc.save(merged_buffer)
    merged_buffer.seek(0)
    
    # Return the modified file
    return send_file(
        merged_buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='merged.docx'
    )


@app.route('/merge-curriculum-syllabi', methods=['POST'])
def merge_curriculum_syllabi():
    """Accepts multipart form-data with keys 'main' and 'syllabi' (files).
    Merges the two DOCX files using docxcompose.Composer and returns the merged file.
    """
    # Ensure files present
    if 'main' not in request.files:
        return {"error": "Missing 'main' file"}, 400

    main_file = request.files['main']
    syllabi_file = request.files.get('syllabi')

    # Load main document
    try:
        main_doc = Document(io.BytesIO(main_file.read()))
    except Exception as e:
        return {"error": f"Failed to read main document: {str(e)}"}, 400

    composer = Composer(main_doc)

    # If a syllabi file is provided, append it
    if syllabi_file:
        try:
            syll_doc = Document(io.BytesIO(syllabi_file.read()))
            # ensure the appended document starts on a fresh page
            try:
                main_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            except Exception:
                # if adding a page break fails, continue and attempt append anyway
                pass
            composer.append(syll_doc)
        except Exception as e:
            return {"error": f"Failed to read/append syllabi document: {str(e)}"}, 400

    # Save merged document to buffer
    merged_buffer = io.BytesIO()
    try:
        composer.save(merged_buffer)
    except Exception as e:
        return {"error": f"Failed to compose merged document: {str(e)}"}, 500

    merged_buffer.seek(0)
    download_name = request.form.get('downloadName') or 'Curriculum_With_Syllabi.docx'

    return send_file(
        merged_buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=download_name
    )

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    app.run(host='0.0.0.0', port=port, debug=False)
